import os
import io
import json
import re
from datetime import date
import httpx
from fastapi import FastAPI, Request, File, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from bs4 import BeautifulSoup
from groq import Groq

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

# ─── Prompts ──────────────────────────────────────────────────────────────────

REVIEW_PROMPT = """You are ClearCV, an expert AI resume reviewer. You analyze resumes against job descriptions and provide structured, actionable feedback.

Return your analysis in EXACTLY this JSON format (no markdown, no code fences, just raw JSON):

{
  "match_score": <number 0-100>,
  "match_summary": "<one sentence explaining the score>",
  "ats_check": {
    "score": "<Good|Fair|Poor>",
    "issues": ["<issue 1>", "<issue 2>"]
  },
  "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "weaknesses": ["<weakness 1>", "<weakness 2>", "<weakness 3>"],
  "rewrite_suggestions": [
    {"original": "<exact line from resume>", "suggested": "<improved version>", "reason": "<why>"}
  ],
  "missing_keywords": ["<keyword 1>", "<keyword 2>", "<keyword 3>", "<keyword 4>", "<keyword 5>"]
}

Rules: Be specific, reference actual content, match score should reflect real alignment, return ONLY valid JSON."""

COVER_LETTER_PROMPT = """You are ClearCV, an expert cover letter writer. Write tailored, professional cover letters.

Analyse the job description carefully — identify what the employer emphasises most (technical skills, leadership, culture, specific tools). The cover letter MUST mirror those priorities.

If additional professional context is provided (GitHub, portfolio), use specific projects and achievements from it to strengthen the letter.

Return ONLY this JSON (no markdown, no code fences):

{
  "candidate_name": "<full name from resume, or 'Your Name'>",
  "company_name": "<company from job description, or 'the Company'>",
  "role": "<exact job title>",
  "paragraphs": [
    "<Opening: strong hook, name the role and company, genuine enthusiasm based on something specific in the JD>",
    "<Body 1: 2-3 specific achievements from resume/portfolio that directly address top JD requirements. Use numbers/metrics.>",
    "<Body 2: address other key requirements, cultural fit, or why this specific company/role>",
    "<Closing: confident call to action>"
  ]
}

Rules: First person, professional but warm, never generic filler, ground every claim in real data, return ONLY valid JSON."""

UPDATED_CV_PROMPT = """You are ClearCV, an expert CV writer and ATS specialist. Rewrite a candidate's CV to be the strongest possible version for a specific role.

You will receive:
1. The original CV
2. The target job description
3. Analysis results (rewrite suggestions, missing keywords, ATS issues to fix)
4. Optional: GitHub repos, portfolio content, other professional context

Your task:
- Fix ALL ATS issues: use standard section headers, remove special characters (✓ ✦ ■ → etc.), fix invisible spaces
- Apply every rewrite suggestion from the analysis
- Naturally weave in all missing keywords throughout
- If GitHub/portfolio shows a more relevant project for this role, swap it in for a less relevant one
- Quantify achievements with numbers, %, $, timescales wherever possible
- Every bullet starts with a strong past/present action verb
- Use ATS-friendly section titles: Professional Summary, Work Experience, Projects, Education, Technical Skills, Certifications
- NEVER fabricate experience — only restructure and improve what exists

Return ONLY this exact JSON (no markdown, no code fences):

{
  "name": "<candidate full name>",
  "contact": {
    "email": "<email>",
    "phone": "<phone>",
    "location": "<city, country>",
    "links": ["<github url>", "<linkedin url>", "<portfolio url>"]
  },
  "sections": [
    {
      "title": "<section title>",
      "body": ["<line 1>", "<line 2>", ...]
    }
  ]
}

Body line formatting:
- Job/project/education headers: "**Role or Degree** | Organisation | Date Range"
- Bullet points: "• Achievement starting with action verb"
- Use "" (empty string) for spacing between items within a section
- Plain text for paragraph sections like Professional Summary
- Skills: comma-separated, grouped by category on one or two lines

Return ONLY valid JSON, no other text."""

# ─── Helpers ──────────────────────────────────────────────────────────────────

async def fetch_github_context(url: str) -> str:
    """Fetch GitHub profile repos via the public GitHub API."""
    try:
        username = url.rstrip("/").split("github.com/")[-1].split("/")[0].lstrip("@")
        if not username:
            return ""
        headers = {"Accept": "application/vnd.github.v3+json"}
        async with httpx.AsyncClient(timeout=10) as http:
            user_res  = await http.get(f"https://api.github.com/users/{username}", headers=headers)
            repos_res = await http.get(
                f"https://api.github.com/users/{username}/repos?sort=updated&per_page=20&type=public",
                headers=headers,
            )
        if user_res.status_code != 200:
            return ""
        user  = user_res.json()
        repos = repos_res.json() if repos_res.status_code == 200 else []
        lines = [f"GitHub: @{username}"]
        if user.get("bio"):
            lines.append(f"Bio: {user['bio']}")
        own = [r for r in repos if isinstance(r, dict) and not r.get("fork")][:12]
        if own:
            lines.append("\nRepositories:")
            for r in own:
                entry = f"- {r.get('name','')}"
                if r.get("language"):   entry += f" [{r['language']}]"
                if r.get("description"): entry += f": {r['description']}"
                if r.get("stargazers_count", 0): entry += f" ({r['stargazers_count']}★)"
                lines.append(entry)
        return "\n".join(lines)
    except Exception:
        return ""


async def fetch_url_context(url: str, label: str = "Link") -> str:
    """Fetch and extract text from any URL."""
    try:
        if not url.startswith("http"):
            url = "https://" + url
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        async with httpx.AsyncClient(timeout=10, follow_redirects=True) as http:
            res = await http.get(url, headers=headers)
            res.raise_for_status()
        soup = BeautifulSoup(res.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        main = soup.find("main") or soup.find("article") or soup.body
        if main:
            lines = [l.strip() for l in main.get_text("\n").split("\n") if l.strip()]
            return "\n".join(lines[:120])
        return ""
    except Exception:
        return ""


async def build_links_context(links: list) -> str:
    """Fetch all provided links and combine into a single context string."""
    contexts = []
    for link in links:
        url   = link.get("url", "").strip()
        label = link.get("label", "Additional Link")
        if not url:
            continue
        if "github.com" in url:
            ctx = await fetch_github_context(url)
            if ctx:
                contexts.append(f"=== GitHub Data ===\n{ctx}")
        else:
            ctx = await fetch_url_context(url, label)
            if ctx:
                contexts.append(f"=== {label} ===\n{ctx[:2000]}")
    return "\n\n".join(contexts)


def build_cv_docx(cv_data: dict) -> io.BytesIO:
    """Generate a professional CV as a .docx file."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)

    ACCENT = RGBColor(0xa7, 0x8b, 0xfa)
    GRAY   = RGBColor(0x50, 0x50, 0x50)

    def add_bottom_border(para, color="a78bfa", size=4):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    str(size))
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def render_bold_line(para, line, size=10):
        parts = re.split(r"\*\*", line)
        for i, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                run.bold = (i % 2 == 1)
                run.font.size = Pt(size)

    # Name
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_after = Pt(3)
    nr = np_.add_run(cv_data.get("name", ""))
    nr.bold = True
    nr.font.size = Pt(18)

    # Contact
    contact = cv_data.get("contact", {})
    parts = []
    for f in ["email", "phone", "location"]:
        if contact.get(f): parts.append(contact[f])
    for lnk in contact.get("links", []):
        if lnk: parts.append(lnk)
    if parts:
        cp = doc.add_paragraph(" · ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY

    # Divider
    div = doc.add_paragraph()
    div.paragraph_format.space_after = Pt(4)
    add_bottom_border(div, color="a78bfa", size=8)

    # Sections
    for section in cv_data.get("sections", []):
        title = section.get("title", "").upper()
        body  = section.get("body", [])

        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(10)
        hdr.paragraph_format.space_after  = Pt(3)
        hr = hdr.add_run(title)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = ACCENT
        add_bottom_border(hdr, color="a78bfa", size=4)

        for line in body:
            if not line:
                ep = doc.add_paragraph()
                ep.paragraph_format.space_after = Pt(2)
                continue
            if line.startswith("•"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Inches(0.15)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(line)
                run.font.size = Pt(10)
            elif "**" in line:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after  = Pt(2)
                render_bold_line(p, line)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(line)
                run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_cv_pdf(cv_data: dict) -> io.BytesIO:
    """Generate a professional CV as a PDF file."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf    = io.BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=letter,
                               rightMargin=inch, leftMargin=inch,
                               topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    ACCENT = colors.HexColor("#a78bfa")
    GRAY   = colors.HexColor("#505050")

    name_s    = ParagraphStyle("N",  fontSize=18, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)
    contact_s = ParagraphStyle("C",  fontSize=9,  fontName="Helvetica",      alignment=TA_CENTER, textColor=GRAY, spaceAfter=8)
    header_s  = ParagraphStyle("H",  fontSize=10, fontName="Helvetica-Bold", textColor=ACCENT, spaceBefore=10, spaceAfter=3)
    body_s    = ParagraphStyle("B",  fontSize=10, fontName="Helvetica",      leading=14, spaceAfter=2)
    bullet_s  = ParagraphStyle("BL", fontSize=10, fontName="Helvetica",      leading=14, spaceAfter=2, leftIndent=12)
    role_s    = ParagraphStyle("R",  fontSize=10, fontName="Helvetica",      leading=14, spaceAfter=2, spaceBefore=5)

    def parse_bold(line):
        return re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", line)

    story = [Paragraph(cv_data.get("name", ""), name_s)]

    contact = cv_data.get("contact", {})
    parts = []
    for f in ["email", "phone", "location"]:
        if contact.get(f): parts.append(contact[f])
    for lnk in contact.get("links", []):
        if lnk: parts.append(lnk)
    if parts:
        story.append(Paragraph(" · ".join(parts), contact_s))

    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceAfter=6))

    for section in cv_data.get("sections", []):
        title = section.get("title", "").upper()
        body  = section.get("body", [])
        story.append(Paragraph(title, header_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT, spaceAfter=4))
        for line in body:
            if not line:
                story.append(Spacer(1, 4))
            elif line.startswith("•"):
                story.append(Paragraph(line, bullet_s))
            elif "**" in line:
                story.append(Paragraph(parse_bold(line), role_s))
            else:
                story.append(Paragraph(line, body_s))

    doc.build(story)
    buf.seek(0)
    return buf


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    with open("index.html", "r") as f:
        return f.read()


@app.post("/api/parse-cv")
async def parse_cv(file: UploadFile = File(...)):
    content  = await file.read()
    filename = file.filename.lower()
    try:
        if filename.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text   = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif filename.endswith(".docx"):
            from docx import Document
            doc  = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        elif filename.endswith(".txt"):
            text = content.decode("utf-8")
        else:
            return JSONResponse({"error": "Unsupported file type. Please upload PDF, DOCX, or TXT."}, status_code=400)
        text = text.strip()
        if not text:
            return JSONResponse({"error": "Could not extract text. Try pasting manually."}, status_code=400)
        return {"text": text}
    except Exception as e:
        return JSONResponse({"error": f"Failed to parse file: {str(e)}"}, status_code=500)


@app.post("/api/extract-url")
async def extract_url(request: Request):
    body = await request.json()
    url  = body.get("url", "").strip()
    if not url:
        return JSONResponse({"error": "URL is required."}, status_code=400)
    if not url.startswith("http"):
        url = "https://" + url
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"}
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as http:
            response = await http.get(url, headers=headers)
            response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()
        selectors = [
            {"id": "job-description"},
            {"class": re.compile(r"job.?description|jobDescription|job.?detail|jobDetail", re.I)},
            {"class": re.compile(r"description__text|job-view-layout|show-more-less-html", re.I)},
        ]
        container = None
        for sel in selectors:
            container = soup.find(attrs=sel)
            if container: break
        if not container:
            container = soup.find("main") or soup.find("article") or soup.body
        if container:
            lines = [l.strip() for l in container.get_text("\n").split("\n") if l.strip()]
            text  = "\n".join(lines)
        else:
            text = ""
        if not text or len(text) < 100:
            return JSONResponse({"error": "Could not extract job description. Try pasting manually."}, status_code=400)
        return {"text": text[:8000]}
    except httpx.HTTPStatusError as e:
        return JSONResponse({"error": f"HTTP {e.response.status_code} — try pasting manually."}, status_code=400)
    except Exception as e:
        return JSONResponse({"error": f"Failed to fetch URL: {str(e)}"}, status_code=500)


@app.post("/api/review")
async def review_resume(request: Request):
    body             = await request.json()
    resume           = body.get("resume", "")
    job_description  = body.get("job_description", "")
    if not resume or not job_description:
        return JSONResponse({"error": "Both inputs are required."}, status_code=400)

    user_message = f"Please review this resume against the job description.\n\n---RESUME---\n{resume}\n\n---JOB DESCRIPTION---\n{job_description}"

    async def generate():
        stream = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": REVIEW_PROMPT}, {"role": "user", "content": user_message}],
            max_tokens=4096, stream=True,
        )
        for chunk in stream:
            content = chunk.choices[0].delta.content
            if content:
                yield f"data: {json.dumps({'text': content})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/api/cover-letter")
async def generate_cover_letter(request: Request):
    body            = await request.json()
    resume          = body.get("resume", "")
    job_description = body.get("job_description", "")
    links           = body.get("links", [])
    if not resume or not job_description:
        return JSONResponse({"error": "Resume and job description are required."}, status_code=400)

    links_context = await build_links_context(links) if links else ""
    user_message  = f"Write a tailored cover letter.\n\n---RESUME---\n{resume}\n\n---JOB DESCRIPTION---\n{job_description}"
    if links_context:
        user_message += f"\n\n---ADDITIONAL PROFESSIONAL CONTEXT---\n{links_context}"

    async def generate():
        stream = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": COVER_LETTER_PROMPT}, {"role": "user", "content": user_message}],
            max_tokens=2048, stream=True,
        )
        for chunk in stream:
            content = chunk.choices[0].delta.content
            if content:
                yield f"data: {json.dumps({'text': content})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/api/updated-cv")
async def generate_updated_cv(request: Request):
    body            = await request.json()
    resume          = body.get("resume", "")
    job_description = body.get("job_description", "")
    analysis        = body.get("analysis", {})
    links           = body.get("links", [])
    if not resume or not job_description:
        return JSONResponse({"error": "Resume and job description are required."}, status_code=400)

    links_context = await build_links_context(links) if links else ""

    # Build analysis summary for the prompt
    analysis_text = ""
    if analysis:
        rewrites   = analysis.get("rewrite_suggestions", [])
        keywords   = analysis.get("missing_keywords", [])
        ats_issues = analysis.get("ats_check", {}).get("issues", [])
        if rewrites:
            analysis_text += "REWRITE SUGGESTIONS:\n"
            for r in rewrites:
                analysis_text += f"- Change: '{r.get('original','')}'\n  To: '{r.get('suggested','')}'\n  Why: {r.get('reason','')}\n"
        if keywords:
            analysis_text += f"\nMISSING KEYWORDS TO INCORPORATE: {', '.join(keywords)}\n"
        if ats_issues:
            analysis_text += "\nATS ISSUES TO FIX:\n" + "\n".join(f"- {i}" for i in ats_issues)

    user_message = f"Rewrite this CV for the target role.\n\n---ORIGINAL CV---\n{resume}\n\n---JOB DESCRIPTION---\n{job_description}"
    if analysis_text:
        user_message += f"\n\n---ANALYSIS (apply all these changes)---\n{analysis_text}"
    if links_context:
        user_message += f"\n\n---ADDITIONAL PROFESSIONAL CONTEXT---\n{links_context}"

    async def generate():
        stream = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": UPDATED_CV_PROMPT}, {"role": "user", "content": user_message}],
            max_tokens=4096, stream=True,
        )
        for chunk in stream:
            content = chunk.choices[0].delta.content
            if content:
                yield f"data: {json.dumps({'text': content})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


# ─── Cover Letter Downloads ────────────────────────────────────────────────────

@app.post("/api/download/pdf")
async def download_cover_letter_pdf(request: Request):
    body = await request.json()
    cl   = body.get("cover_letter", {})

    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            rightMargin=1.25*inch, leftMargin=1.25*inch,
                            topMargin=inch, bottomMargin=inch)

    styles     = getSampleStyleSheet()
    name_s     = ParagraphStyle("Name",  parent=styles["Normal"], fontSize=15, fontName="Helvetica-Bold", spaceAfter=2)
    date_s     = ParagraphStyle("Date",  parent=styles["Normal"], fontSize=11, spaceAfter=14)
    open_s     = ParagraphStyle("Open",  parent=styles["Normal"], fontSize=11, spaceAfter=14)
    body_s     = ParagraphStyle("Body",  parent=styles["Normal"], fontSize=11, leading=17, spaceAfter=12, alignment=TA_LEFT)
    close_s    = ParagraphStyle("Close", parent=styles["Normal"], fontSize=11, spaceBefore=14)

    story = [
        Paragraph(cl.get("candidate_name", ""), name_s),
        Paragraph(date.today().strftime("%B %d, %Y"), date_s),
        Paragraph(f"Re: {cl.get('role','Application')} — {cl.get('company_name','')}", date_s),
        Spacer(1, 6),
        Paragraph("Dear Hiring Manager,", open_s),
    ]
    for para in cl.get("paragraphs", []):
        story.append(Paragraph(para, body_s))
    story += [Spacer(1, 8), Paragraph("Sincerely,", close_s), Spacer(1, 36), Paragraph(cl.get("candidate_name", ""), close_s)]

    doc.build(story)
    buf.seek(0)
    fname = f"Cover_Letter_{cl.get('company_name','Application').replace(' ','_')}.pdf"
    return Response(content=buf.read(), media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@app.post("/api/download/docx")
async def download_cover_letter_docx(request: Request):
    body = await request.json()
    cl   = body.get("cover_letter", {})

    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    def para(text, bold=False, size=11, after=12):
        p    = doc.add_paragraph()
        run  = p.add_run(text)
        run.bold       = bold
        run.font.size  = Pt(size)
        p.paragraph_format.space_after = Pt(after)
        return p

    para(cl.get("candidate_name",""), bold=True, size=15, after=2)
    para(date.today().strftime("%B %d, %Y"), after=2)
    para(f"Re: {cl.get('role','Application')} — {cl.get('company_name','')}", after=14)
    para("Dear Hiring Manager,", after=12)
    for p_text in cl.get("paragraphs", []):
        para(p_text, after=12)
    doc.add_paragraph()
    para("Sincerely,", after=0)
    doc.add_paragraph()
    doc.add_paragraph()
    para(cl.get("candidate_name",""), after=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"Cover_Letter_{cl.get('company_name','Application').replace(' ','_')}.docx"
    return Response(content=buf.read(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


# ─── Updated CV Downloads ─────────────────────────────────────────────────────

@app.post("/api/download/updated-cv-pdf")
async def download_updated_cv_pdf(request: Request):
    body    = await request.json()
    cv_data = body.get("cv_data", {})
    buf     = build_cv_pdf(cv_data)
    name    = cv_data.get("name", "CV").replace(" ", "_")
    return Response(content=buf.read(), media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="Updated_CV_{name}.pdf"'})


@app.post("/api/download/updated-cv-docx")
async def download_updated_cv_docx(request: Request):
    body    = await request.json()
    cv_data = body.get("cv_data", {})
    buf     = build_cv_docx(cv_data)
    name    = cv_data.get("name", "CV").replace(" ", "_")
    return Response(content=buf.read(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="Updated_CV_{name}.docx"'})


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
